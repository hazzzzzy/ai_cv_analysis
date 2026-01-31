from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Tuple
import uuid

from docx import Document
from pypdf import PdfReader

from app.core.config import settings


def parse_resume_content(filename: str, content: bytes) -> Tuple[str, dict]:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _parse_pdf(content, filename)
    if name.endswith(".docx"):
        return _parse_docx(content, filename)
    return _parse_unknown(content, filename)


def _parse_pdf(content: bytes, filename: str) -> Tuple[str, dict]:
    try:
        reader = PdfReader(BytesIO(content))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts).strip()
        meta = {
            "parser": "pypdf",
            "filename": filename,
            "size_bytes": len(content),
            "page_count": len(reader.pages),
        }
        print(text)
        return text, meta
    except Exception:
        return "", {
            "parser": "pypdf",
            "filename": filename,
            "size_bytes": len(content),
            "error": "pdf 解析失败",
        }


def _parse_docx(content: bytes, filename: str) -> Tuple[str, dict]:
    try:
        doc = Document(BytesIO(content))
        parts = [para.text for para in doc.paragraphs if para.text]
        text = "\n".join(parts).strip()
        meta = {
            "parser": "python-docx",
            "filename": filename,
            "size_bytes": len(content),
            "paragraph_count": len(doc.paragraphs),
        }
        return text, meta
    except Exception:
        return "", {
            "parser": "python-docx",
            "filename": filename,
            "size_bytes": len(content),
            "error": "docx 解析失败",
        }


def _parse_unknown(content: bytes, filename: str) -> Tuple[str, dict]:
    meta = {
        "parser": "unknown",
        "filename": filename,
        "size_bytes": len(content),
    }
    return "", meta


def save_resume_file(filename: str, content: bytes) -> str:
    base_dir = Path(settings.resume_local_dir)
    base_dir.mkdir(parents=True, exist_ok=True)
    suffix = Path(filename).suffix if filename else ""
    safe_name = f"{uuid.uuid4().hex}{suffix}"
    path = base_dir / safe_name
    path.write_bytes(content)
    return str(path)
